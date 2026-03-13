from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def force_run_formatting(run, font_name='Sora', font_size_pt=12, color_rgb=None, bold=False, all_caps=False):
    font = run.font
    font.name = font_name
    font.size = Pt(font_size_pt)
    if color_rgb:
        font.color.rgb = color_rgb
    font.bold = bold
    font.all_caps = all_caps
    
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def add_bottom_border(paragraph_style, color_hex='9FEC14', size="6", space="1"):
    """
    Adiciona borda inferior ao estilo de parágrafo via XML.
    """
    pPr = paragraph_style.element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        pBdr.append(bottom)
        
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size) # 1/8 pt units. 6 = 0.75pt, 12 = 1.5pt
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color_hex)

def setup_styles(doc):
    NEON_GREEN = RGBColor(*hex_to_rgb('9FEC14'))
    STD_TEXT = RGBColor(*hex_to_rgb('595959'))

    # Styles
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Sora'
    font_normal.size = Pt(12)
    font_normal.color.rgb = STD_TEXT
    
    pf_normal = style_normal.paragraph_format
    pf_normal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf_normal.first_line_indent = Cm(3)
    pf_normal.space_after = Pt(12)
    pf_normal.line_spacing = 1.5

    # Criar estilo para Header (Logo)
    # Não precisa estilo especifico, vamos inserir direto.

    # Criar estilo "GreenLine" (Divisor)
    if 'GreenLine' not in doc.styles:
        style_gl = doc.styles.add_style('GreenLine', 1) # 1 = Paragraph Style
        pf_gl = style_gl.paragraph_format
        pf_gl.space_before = Pt(0)
        pf_gl.space_after = Pt(12) # Espaço depois da linha
        pf_gl.line_spacing = 1
        # Adicionar Borda
        add_bottom_border(style_gl, color_hex='9FEC14', size="12") # Borda mais grossa

def create_master_template():
    print("🎨 Criando MASTER TEMPLATE V8 (Logo + Green Lines)...")
    
    doc = Document()
    setup_styles(doc)
    
    NEON_GREEN = RGBColor(*hex_to_rgb('9FEC14'))
    STD_TEXT = RGBColor(*hex_to_rgb('595959'))
    GRAY_TEXT = RGBColor(*hex_to_rgb('656565'))
    
    # --- HEADER COM LOGO ---
    # Tentar inserir logo se existir
    logo_path = 'assets/imagens-fixas/logo-genesis.png'
    if os.path.exists(logo_path):
        section = doc.sections[0]
        header = section.header
        p_head = header.paragraphs[0]
        p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_head.add_run()
        run_img.add_picture(logo_path, width=Cm(4)) # Tamanho ajustado
    else:
        print(f"⚠️ Aviso: Logo não encontrada em {logo_path}")

    # --- BODY (CAPA) ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('LAUDO DE VIABILIDADE')
    force_run_formatting(run, font_size_pt=26, color_rgb=NEON_GREEN, bold=True, all_caps=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    force_run_formatting(p.add_run('MARCA: '), font_size_pt=14, color_rgb=GRAY_TEXT, bold=True)
    force_run_formatting(p.add_run('{{ marca_principal }}\n'), font_size_pt=14, color_rgb=GRAY_TEXT)
    force_run_formatting(p.add_run('CLIENTE: '), font_size_pt=14, color_rgb=GRAY_TEXT, bold=True)
    force_run_formatting(p.add_run('{{ cliente }}\n'), font_size_pt=14, color_rgb=GRAY_TEXT)
    force_run_formatting(p.add_run('DATA: '), font_size_pt=14, color_rgb=GRAY_TEXT, bold=True)
    force_run_formatting(p.add_run('{{ data_analise }}'), font_size_pt=14, color_rgb=GRAY_TEXT)
    
    doc.add_page_break()
    
    # --- LOOP HELPER ---
    def add_loop_section_with_line(variable_name):
        # Start Loop
        p_start = doc.add_paragraph()
        p_start.paragraph_format.space_after = Pt(0)
        run_start = p_start.add_run(f"{{% for line in {variable_name} %}}")
        force_run_formatting(run_start, font_size_pt=1, color_rgb=RGBColor(255,255,255)) # Invisível se possível

        # Content Paragraph
        p_content = doc.add_paragraph()
        p_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_content.paragraph_format.first_line_indent = Cm(3)
        p_content.paragraph_format.space_after = Pt(6) # Espaço menor antes da linha
        p_content.paragraph_format.line_spacing = 1.5
        run_content = p_content.add_run("{{ line }}")
        force_run_formatting(run_content, font_name='Sora', font_size_pt=12, color_rgb=STD_TEXT)

        # Green Line Separator
        # Criamos um parágrafo vazio com o estilo 'GreenLine' que tem borda inferior
        p_line = doc.add_paragraph(style='GreenLine')
        # Para garantir que a linha apareça, precisamos de algum conteúdo ou o Word ignora?
        # Às vezes espaço em branco basta. Vamos testar.
        # Se não aparecer, adicionar um espaço.
        # force_run_formatting(p_line.add_run(" "), font_size_pt=1) 
        
        # End Loop
        p_end = doc.add_paragraph()
        p_end.paragraph_format.space_after = Pt(0)
        run_end = p_end.add_run("{% endfor %}")
        force_run_formatting(run_end, font_size_pt=1, color_rgb=RGBColor(255,255,255))

    # --- TITLES ---
    def add_section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        force_run_formatting(run, font_size_pt=14, color_rgb=NEON_GREEN, bold=True, all_caps=True)

    def add_subsection_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        force_run_formatting(run, font_size_pt=12, color_rgb=NEON_GREEN, bold=True, all_caps=True)

    # --- DOCUMENT STRUCTURE ---
    add_section_title('1. CLASSES RECOMENDADAS')
    add_loop_section_with_line('classes_sugeridas')
    
    add_section_title('2. ANÁLISE DE REQUISITOS')
    
    add_subsection_title('VERACIDADE')
    add_loop_section_with_line('veracidade')
    
    add_subsection_title('LICITUDE')
    add_loop_section_with_line('licitude')
    
    add_subsection_title('DISTINTIVIDADE')
    add_loop_section_with_line('distintividade')
    
    add_section_title('3. ANÁLISE DE COLIDÊNCIAS')
    add_loop_section_with_line('analises_colidencias')
    
    add_section_title('4. CONCLUSÃO E ESTRATÉGIA')
    add_loop_section_with_line('conclusao_final')
    
    # Assinatura
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run('É o parecer.\n\nEquipe Vialum')
    force_run_formatting(r, font_size_pt=12, color_rgb=STD_TEXT)
    
    # Rodapé
    section = doc.sections[0]
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.first_line_indent = Cm(0)
    r = p_foot.add_run("Vialum - Inteligência em Marcas | Confidencial")
    force_run_formatting(r, font_size_pt=8, color_rgb=GRAY_TEXT)

    output_path = 'templates/master-template.docx'
    doc.save(output_path)
    print(f"✅ Template Master V8 (Header Logo + Green Lines) salvo em: {output_path}")

if __name__ == "__main__":
    create_master_template()
