import re
import os
import sys
import shutil
from pathlib import Path
from datetime import datetime
import time

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google_drive_service

# --- CONFIGURAÇÕES VISUAIS ---
COLOR_PRIMARY = RGBColor(159, 236, 20) # #9FEC14
COLOR_TEXT = RGBColor(89, 89, 89)      # #595959
COLOR_GRAY = RGBColor(101, 101, 101)   # #656565

# --- UTILS ---
def set_margins(doc):
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(4) 
    section.bottom_margin = Cm(2.5)
    section.different_first_page_header_footer = True

def format_run(run, font_name='Sora', font_size=12, color=COLOR_TEXT, bold=False, all_caps=False, underline=False):
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.color.rgb = color
    font.bold = bold
    font.all_caps = all_caps
    font.underline = underline
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def add_bottom_border(paragraph, color_hex='9FEC14', size="24"):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def data_por_extenso(data_str):
    meses = {
        '01': 'janeiro', '02': 'fevereiro', '03': 'março', '04': 'abril',
        '05': 'maio', '06': 'junho', '07': 'julho', '08': 'agosto',
        '09': 'setembro', '10': 'outubro', '11': 'novembro', '12': 'dezembro'
    }
    try:
        dia, mes, ano = data_str.split('/')
        return f"{dia} de {meses.get(mes, mes)} de {ano}"
    except:
        return data_str

# --- PARSING ---
def extrair_campo(texto: str, padrao: str, default: str = '') -> str:
    match = re.search(padrao, texto)
    return match.group(1).strip() if match else default

def extrair_secao(texto: str, padrao: str, flags=0) -> str:
    match = re.search(padrao, texto, flags=flags)
    return match.group(1).strip() if match else ''

def parsear_plano(caminho_md):
    with open(caminho_md, 'r', encoding='utf-8') as f:
        conteudo = f.read()
    
    dados = {}
    dados['cliente'] = extrair_campo(conteudo, r'\*\*Cliente:\*\*\s*(.+)')
    dados['marca_principal'] = extrair_campo(conteudo, r'\*\*Marca Principal:\*\*\s*(.+)')
    dados['data_analise'] = extrair_campo(conteudo, r'\*\*Data:\*\*\s*(.+)', default=datetime.now().strftime('%d/%m/%Y'))
    
    dados['secoes'] = {}
    
    # 1. Classes: Extrai entre o título e a próxima seção (## PARTE 2 ou ## ANÁLISE DOS REQUISITOS)
    classes = extrair_secao(conteudo, r'## LAUDO DESCRITIVO POR CLASSE\n*(.*?)\n*(?=##|$)', flags=re.DOTALL)
    dados['secoes']['Classes Recomendadas'] = classes
    
    # 2. Requisitos (Narrativa V6): Extrai entre o título e a próxima seção (## LAUDO DESCRITIVO POR CLASSE)
    # Corrigido: aceita qualquer grafia (LICITUDE, LICEIDADE, etc)
    req = extrair_secao(conteudo, r'## ANÁLISE DOS REQUISITOS \(.*?\)\n*(.*?)\n*(?=## LAUDO DESCRITIVO POR CLASSE|$)', flags=re.DOTALL)
    dados['secoes']['Requisitos'] = req
    
    # 3. Colidências e Estratégia (Tudo da Parte 2)
    colid = extrair_secao(conteudo, r'## PARTE 2: ANÁLISE DE COLIDÊNCIAS E ANTERIORIDADES\n*(.*)', flags=re.DOTALL)
    dados['secoes']['Colidências'] = colid
    
    return dados

def create_docx(dados, output_path):
    print(f"🏗️ Construindo DOCX (Builder V4 - Fixes)...")
    doc = Document()
    set_margins(doc)
    section = doc.sections[0]
    
    # Caminhos Absolutos
    base_dir = Path(__file__).parent.parent.absolute()
    logo_path = base_dir / "resources/assets/imagens-fixas/logo-genesis.png"
    fluxo_img = base_dir / "resources/assets/imagens-fixas/fluxograma-inpi.png"

    # --- HEADER ---
    header_first = section.first_page_header
    table = header_first.add_table(rows=1, cols=2, width=Cm(16))
    table.autofit = False
    
    cell_left = table.cell(0, 0)
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path.exists():
       run_l = p_left.add_run()
       run_l.add_picture(str(logo_path), width=Cm(7))
    
    cell_right = table.cell(0, 1)
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_para = p_right.add_run("Para\n")
    format_run(run_para, font_size=11, color=COLOR_GRAY, bold=False)
    run_cli = p_right.add_run(f"{dados['cliente']}")
    format_run(run_cli, font_size=16, color=COLOR_TEXT, bold=True)

    header_std = section.header
    p_std = header_std.paragraphs[0]
    p_std.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path.exists():
        run_std = p_std.add_run()
        run_std.add_picture(str(logo_path), width=Cm(6))
    
    # --- FOOTER ---
    footer_text = "Genesis - Marcas e Patentes | Rua Pará, 945, Londrina, Paraná, CEP 86010-450 | @genesisregistro"
    def fill_footer(footer_obj):
        p = footer_obj.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(footer_text)
        format_run(run, font_size=9, color=COLOR_GRAY)
    fill_footer(section.footer)
    fill_footer(section.first_page_footer)

    # --- HELPERS ---
    def add_para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.first_line_indent = Cm(3) if align == WD_ALIGN_PARAGRAPH.JUSTIFY else 0
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        format_run(run, font_size=12, color=COLOR_TEXT)
        return p

    def add_main_title(text, page_break=True):
        if page_break:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True # Garante que o título não fique só
        run = p.add_run(text)
        format_run(run, font_size=16, color=COLOR_PRIMARY, bold=True, all_caps=True)
        add_bottom_border(p)

    def add_req_inline(title, text):
        # Tópico inline: **TITLE** - Text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # p.paragraph_format.first_line_indent = Cm(3) # Talvez sem indent para destacar? O user mandou bullet com traço.
        # Vamos manter o padrão de texto corrido (recuo 3cm) ou remover recuo para tópicos?
        # User disse "nos mesmos termos dos outros títulos" para INTRODUÇÃO, mas para esses disse "tópicos que deveriam ter em texto corrido"
        # Vou usar indent padrão 3cm para alinhar com o texto.
        p.paragraph_format.first_line_indent = Cm(3)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        
        # TITLE (Bold + Underline)
        r_title = p.add_run(title)
        format_run(r_title, font_size=12, color=COLOR_TEXT, bold=True, underline=True)
        
        # Separator " - "
        r_sep = p.add_run(" - ")
        format_run(r_sep, font_size=12, color=COLOR_TEXT)
        
        # Text
        r_text = p.add_run(text)
        format_run(r_text, font_size=12, color=COLOR_TEXT)

    def add_content_block(raw_text):
        if not raw_text: return
        paragraphs = raw_text.split('\n')
        for text in paragraphs:
            text = text.strip()
            if not text: continue
            
            # FILTRO 1: Remove separadores Markdown (---)
            if re.match(r'^-{3,}$', text): continue
            
            # FILTRO 2: Detectar títulos ## e converter para H2 formatado
            if text.startswith('## '):
                title_text = text[3:].strip()
                title_text = re.sub(r'\*\*(.*?)\*\*', r'\1', title_text)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(title_text)
                format_run(run, font_size=14, color=COLOR_PRIMARY, bold=True)
                add_bottom_border(p)
                continue
            
            # FILTRO 3: Detectar ### e converter para subtítulo
            if text.startswith('### '):
                subtitle_text = text[4:].strip()
                # Remove números do início e negritos
                subtitle_text = re.sub(r'^\d+\.\s*', '', subtitle_text)
                subtitle_text = re.sub(r'\*\*(.*?)\*\*', r'\1', subtitle_text)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(subtitle_text)
                format_run(run, font_size=13, color=COLOR_PRIMARY, bold=True)
                continue
            
            # Detecção de Classes V7.2: **Classe XX (nível)**
            if text.startswith('**Classe'):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Cm(0)
                
                # Processa negrito do Markdown
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        content = part[2:-2]
                        run = p.add_run(content)
                        format_run(run, font_size=13, color=COLOR_PRIMARY, bold=True)
                    else:
                        run = p.add_run(part)
                        format_run(run, font_size=12, color=COLOR_TEXT, bold=False)
                continue
            
            # Processa negrito inline para parágrafos normais
            if '**' in text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(3)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.5
                
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        content = part[2:-2]
                        run = p.add_run(content)
                        format_run(run, font_size=12, color=COLOR_TEXT, bold=True)
                    else:
                        run = p.add_run(part)
                        format_run(run, font_size=12, color=COLOR_TEXT, bold=False)
                continue
            
            add_para(text)


    # --- BODY ---
    
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("LAUDO DE VIABILIDADE")
    format_run(run, font_size=26, color=COLOR_PRIMARY, bold=True, all_caps=True)

    # 1. INTRODUÇÃO
    add_main_title("1. INTRODUÇÃO", page_break=False) 

    marca = dados['marca_principal']
    add_para(f"Este documento tem como objetivo avaliar a viabilidade de registrar a marca denominada “{marca}”, considerando o risco de rejeição. Nossa conclusão é uma avaliação baseada em opinião e não garante que o Instituto Nacional da Propriedade Intelectual (INPI) irá deferir ou negar o pedido. No entanto, garantimos que realizamos todos os esforços necessários na pesquisa de similaridade de marca.")
    add_para("Independentemente do resultado, a decisão de registrar a marca cabe ao cliente. Estaremos focados em apresentar o laudo da melhor forma possível.")
    add_para("Considerando o exposto, é possível efetuar o registro de marcas em diferentes formas, tais como nominativa (baseada no nome), mista (combinando um logo e o nome), figurativa (centrada apenas no logo) e até mesmo tridimensional (representando a configuração física de um produto específico).")
    add_para("Nesta circunstância, após a submissão do requerimento de registro de marca ao INPI, será divulgado na RPI – Revista da Propriedade Industrial, dando início ao prazo legal de 60 (sessenta) dias para que terceiros possam opor-se, sendo concedido ao titular o mesmo período para apresentar sua defesa.")
    add_para("Após esse procedimento, o INPI analisará as oposições apresentadas e realizará a avaliação substancial do pedido de registro da marca, assegurando-se de que esta atenda aos quatro requisitos estabelecidos. Estes requisitos são:")

    # REQUISITOS (Inline Bold+Underline)
    add_req_inline("VERACIDADE", "A marca não pode ter caráter enganoso que induza o público a erro quanto à origem, procedência, natureza, qualidade ou utilidade do produto ou serviço a que se destina. – Previsto no inciso X do art. 124 da LPI.")
    add_req_inline("LICEIDADE", "Esse requisito dispõe que a marca não pode conter elementos característicos presentes em símbolos oficiais, e também não pode conter elementos que contrariam à moral e aos bons costumes. – Previsto nos incisos I, III, XI e XIV do art. 124 da LPI.")
    add_req_inline("DISTINTIVIDADE", "Esse requisitos dispõe que é condição fundamental e função da marca enquanto signo diferenciador de produtos e serviços. – Previsto nos incisos incisos II, VI, VII, VIII, XVIII e XXI do art. 124 da LPI.")
    add_req_inline("NOVIDADE / DISPONIBILIDADE", "Os aspectos gráfico, fonético e ideológico da marca precisam ser exibido de uma forma diferente dos demais já existentes em determinado ramo ou classe. – Previsto nos arts. 124, incisos IV, V, IX, XII, XIII, XV, XVI, XVII, XIX, XX, XXII e XXIII, 125 e 126 da LPI.")

    add_para("Entre os requisitos mencionados, a partir da análise de vários pedidos que foram indeferidos, nota-se que o critério da novidade é frequentemente o principal motivo de indeferimento, devido à sua natureza abrangente e complexa.")
    add_para("Após avaliar o pedido à luz dos quatro requisitos estipulados, o INPI emitirá uma decisão, que poderá ser de indeferimento, concedendo à parte um prazo de 60 dias para apresentar recurso, ou de deferimento, ocasião em que será emitida a certidão de concessão da marca dentro do mesmo prazo após a decisão.")
    add_para("Com a concessão da marca, o INPI procederá com uma nova publicação na RPI – Revista da Propriedade Industrial, marcando o início do prazo de 180 dias (cento e oitenta) dias para que terceiros possam apresentar Processo Administrativo de Nulidade, buscando anular a marca. No caso de tal pedido de nulidade, o titular da marca terá o prazo de 60 dias para contestá-lo, sendo a decisão final sobre o assunto de competência do Presidente do INPI.")
    add_para("Para facilitar melhor a compreensão do processo de registro de marca, anexa-se o fluxograma abaixo demonstrando expressamento o processo adotado pelo INPI até concessão do registro da marca:")

    # FLUXOGRAMA (Abs Path)
    base_dir = Path(__file__).parent.parent.absolute()
    fluxo_img = base_dir / "resources/assets/imagens-fixas/fluxograma-inpi.png"

    if fluxo_img.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(str(fluxo_img), width=Cm(15))
    else:
        print(f"⚠️ Imagem não encontrada: {fluxo_img}")
    
    add_para("Diante do exposto introdutório sobre o procedimento adotado pelo INPI para registro de marca, prosseguimos agora com uma análise voltada ao segmento de mercado do cliente. Nessa análise, buscamos identificar as classes mais adequadas para registro, visando ampliar a proteção da marca em todo o seu ramo.")

    # --- SEÇÕES RENÚMERADAS ---
    add_main_title("2. CLASSES RECOMENDADAS")
    add_content_block(dados['secoes']['Classes Recomendadas'])
    
    add_main_title("3. ANÁLISE DE REQUISITOS (ESPECÍFICA)") 
    add_content_block(dados['secoes']['Requisitos'])
    
    add_main_title("4. ANÁLISE DE COLIDÊNCIAS E VIABILIDADE")
    add_content_block(dados['secoes']['Colidências'])

    # Sign
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(36)
    data_ext = data_por_extenso(dados['data_analise'])
    sign_text = f"É o parecer.\n{data_ext}.\nEquipe Genesis"
    run = p.add_run(sign_text)
    format_run(run, font_size=12, color=COLOR_TEXT)

    doc.save(output_path)
    print(f"✅ DOCX Builder Salvo: {output_path}")

    # --- UPLOAD GOOGLE DRIVE ---
    try:
        google_drive_service.upload_file(str(output_path), marca, dados.get('cliente'))
    except Exception as e:
        print(f"⚠️ Erro ao tentar upload para o Google Drive: {e}")

def main():
    start_time = time.time()
    if len(sys.argv) < 2:
        print("Uso: python gerar_docx_builder.py <caminho_plano_md> [caminho_saida_docx]")
        return
        
    plano_path = Path(sys.argv[1]).absolute()
    if not plano_path.exists():
        print(f"❌ Erro: Arquivo não encontrado: {plano_path}")
        return

    dados = parsear_plano(plano_path)
    marca = dados.get('marca_principal', 'Empresa')

    if len(sys.argv) >= 3:
        output = Path(sys.argv[2]).absolute()
    else:
        output = plano_path.parent / f"{marca} - LAUDO DE VIABILIDADE.docx"
        
    # --- ARQUIVAMENTO LOCAL ---
    # Move todos os arquivos anteriores que comecem com o prefixo da marca e terminem com extensões de saída
    arquivados_dir = plano_path.parent / "arquivados"
    prefixo_busca = f"{marca} -"
    extensoes_alvo = ['.docx']  # Só arquiva DOCX antigo — não toca no PDF

    # Busca arquivos que correspondam ao padrão, mas exclui o próprio plano de análise fonte
    try:
        if not arquivados_dir.exists():
            arquivados_dir.mkdir(parents=True)

        for file_path in plano_path.parent.iterdir():
            if file_path.is_file() and file_path.name.startswith(prefixo_busca):
                # Não arquivar o próprio plano de análise fonte que estamos lendo agora
                if file_path.name == plano_path.name:
                    continue
                # Arquivar apenas DOCX antigos deste script
                if any(file_path.name.endswith(ext) for ext in extensoes_alvo):
                    dest_path = arquivados_dir / file_path.name
                    print(f"📦 Arquivando versão local antiga: {file_path.name}")
                    shutil.move(str(file_path), str(dest_path))
    except Exception as e:
        print(f"⚠️ Erro ao arquivar arquivos locais: {e}")

    create_docx(dados, output)

    # --- GERAR METADADOS "SOBRE O LAUDO" ---
    duration = time.time() - start_time
    duration_str = f"{duration:.2f} segundos"
    
    # Extrair conclusão final (parágrafo final da seção Colidências)
    colid_text = dados['secoes'].get('Colidências', '')
    conclusao = "Não extraída."
    if colid_text:
        links = colid_text.split('\n')
        for line in reversed(links):
            if line.strip() and len(line.strip()) > 50:
                conclusao = line.strip()
                break

    meta_content = f"""# {marca} - Sobre o laudo

**Data do Laudo:** {dados['data_analise']}
**Tempo de Geração:** {duration_str}
**Dono do Laudo (Cliente):** {dados['cliente']}
**Marca Principal:** {marca}

## Conclusão Final
{conclusao}

## Log de Alterações
- [x] Renomeação automática de documentos para padrão "[Marca] - LAUDO DE VIABILIDADE.docx".
- [x] Integração de metadados automáticos via arquivo .md.
- [x] Timming de performance de geração habilitado.
- [x] Upload automático para Google Drive (Possiveis Clientes).
"""
    meta_file_path = plano_path.parent / f"{marca} - Sobre o laudo.md"
    with open(meta_file_path, 'w', encoding='utf-8') as f:
        f.write(meta_content)
    print(f"📝 Arquivo de metadados gerado: {meta_file_path}")

    # Upload do arquivo de metadados para o Drive
    try:
        google_drive_service.upload_file(str(meta_file_path), marca, dados.get('cliente'))
    except Exception as e:
        print(f"⚠️ Erro ao tentar upload do metadata: {e}")

if __name__ == "__main__":
    main()
