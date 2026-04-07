#!/usr/bin/env python3
"""
Assembles Contract and Power of Attorney PDFs using ReportLab.
Usage: python3 assemble_contract.py '<json_data>'

Required JSON keys (PF):
  nome, nacionalidade, estado_civil, profissao, rg, cpf, endereco, cep,
  nome_marca, qtd_classes, valor_total, forma_pagamento, condicao_pagamento

Required JSON keys (PJ) — additionally:
  tipo_pessoa="PJ", razao_social, cnpj, endereco_sede, cep_sede,
  nome_rep, cpf_rep, rg_rep, cargo_rep, nacionalidade_rep,
  estado_civil_rep, profissao_rep, endereco_rep, orgao_emissor_rep

Optional:
  lista_classes, num_parcelas, nome_arquivo_suffix,
  gerar_contrato (bool), gerar_procuracao (bool), orgao_emissor
"""

import sys, os, json, datetime, re
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
AGENT_DIR   = os.path.dirname(SCRIPT_DIR)                  # .agent
ASSETS_DIR  = os.path.join(AGENT_DIR, 'resources', 'assets')
TERMOS_DIR  = os.path.join(AGENT_DIR, 'resources', 'termos')
LOGO_PATH   = os.path.join(ASSETS_DIR, 'logo-genesis.png')

# ── Clause Library (hardcoded for reliability) ─────────────────────────
CLAUSES = {
    "classes": {
        "quantidade": 'prestar serviços relativos ao registro de marca \u201c{{NOME_MARCA}}\u201d junto ao Instituto Nacional de Propriedade Intelectual, em **{{QTD_CLASSES}} classes**, da escolha do CONTRATANTE',
        "especifica": 'prestar serviços relativos ao registro de marca \u201c{{NOME_MARCA}}\u201d junto ao Instituto Nacional de Propriedade Intelectual, especificamente nas classes **{{LISTA_CLASSES}}**, até decisão final',
    },
    "pagamento": {
        "cartao": 'O pagamento será realizado na modalidade **Cartão de Crédito**, sendo disponibilizado link de pagamento para o pagamento das taxas.\n\nO CONTRATANTE compromete-se a efetuar ao CONTRATADO o pagamento da importância de **{{VALOR_TOTAL}}**, que serão integralmente pagos via link de pagamento.',
        "pix_manual": 'O pagamento será realizado na modalidade **PIX (Transferência Manual)**.\n**Dados para Pagamento:**\n- **Chave PIX (CNPJ):** 51.829.412/0001-70\n- **Favorecido:** LUAN VIEIRA MENDES LTDA\n\nO CONTRATANTE compromete-se a efetuar ao CONTRATADO o pagamento da importância de **{{VALOR_TOTAL}}**, mediante envio do comprovante ao CONTRATADO.',
        "pix_auto": 'O pagamento será realizado na modalidade **PIX (Automático)**, sendo disponibilizado Código BR Code (Copia e Cola) ou QR Code gerado pelo sistema de pagamento para a quitação.\n\nO CONTRATANTE compromete-se a efetuar ao CONTRATADO o pagamento da importância de **{{VALOR_TOTAL}}**, através do código fornecido.',
        "boleto_manual": 'O pagamento será realizado na modalidade **Boleto Bancário**, sendo o documento enviado em formato PDF pelo CONTRATADO ao CONTRATANTE.\n\nO CONTRATANTE compromete-se a efetuar ao CONTRATADO o pagamento da importância de **{{VALOR_TOTAL}}**, através da quitação do boleto enviado.',
        "boleto_auto": 'O pagamento será realizado na modalidade **Boleto Bancário**, sendo disponibilizado link para acesso e impressão do documento através do sistema de pagamento.\n\nO CONTRATANTE compromete-se a efetuar ao CONTRATADO o pagamento da importância de **{{VALOR_TOTAL}}**, através do link fornecido.',
    },
    "condicoes": {
        "vista": 'O valor acima descrito deverá ser quitado **à vista**, no ato da contratação.',
        "parcelado": 'O valor acima descrito poderá ser parcelado em até **{{NUM_PARCELAS}} vezes**, sem juros, na opção Cartão de Crédito.',
    },
}

# ── Font Setup ─────────────────────────────────────────────────────────
def setup_fonts():
    fn, fb = "Helvetica", "Helvetica-Bold"
    sora_r = os.path.join(ASSETS_DIR, 'Sora-Regular.ttf')
    sora_b = os.path.join(ASSETS_DIR, 'Sora-Bold.ttf')
    if os.path.exists(sora_r):
        try:
            pdfmetrics.registerFont(TTFont('Sora', sora_r))
            if os.path.exists(sora_b):
                pdfmetrics.registerFont(TTFont('Sora-Bold', sora_b))
            fn, fb = 'Sora', 'Sora-Bold'
        except Exception as e:
            print(f"[WARN] Font error: {e}")
    return fn, fb

# ── Styles ─────────────────────────────────────────────────────────────
def build_styles(fn, fb):
    ss = getSampleStyleSheet()
    # Contract: 12pt, leading 18pt (1.5 spacing)
    ss.add(ParagraphStyle('GN',   fontName=fn, fontSize=12,  leading=18,   alignment=TA_JUSTIFY, firstLineIndent=2*cm, spaceAfter=6))
    ss.add(ParagraphStyle('GNI',  fontName=fn, fontSize=12,  leading=18,   alignment=TA_JUSTIFY, firstLineIndent=0,    spaceAfter=6))
    ss.add(ParagraphStyle('GH1',  fontName=fb, fontSize=16,  leading=22,   alignment=TA_CENTER,  spaceAfter=12))
    ss.add(ParagraphStyle('GH2',  fontName=fb, fontSize=12,  leading=14,   alignment=TA_LEFT,    spaceBefore=8, spaceAfter=2))
    ss.add(ParagraphStyle('GSIG', fontName=fn, fontSize=11,  leading=13,   alignment=TA_CENTER))
    # Procuracao: 11pt, leading 16.5pt
    ss.add(ParagraphStyle('PN',   fontName=fn, fontSize=11,  leading=16.5, alignment=TA_JUSTIFY, firstLineIndent=2*cm, spaceAfter=4))
    ss.add(ParagraphStyle('PNI',  fontName=fn, fontSize=11,  leading=16.5, alignment=TA_JUSTIFY, firstLineIndent=0,    spaceAfter=4))
    return ss

# ── Header/Footer ─────────────────────────────────────────────────────
def _header_footer(canvas, doc):
    canvas.saveState()
    if os.path.exists(LOGO_PATH):
        # Logo is 800x500px → aspect 1.6:1
        # Display at 40mm wide, height ~25mm
        logo_w = 40*mm
        logo_h = 25*mm
        x = (A4[0] - logo_w) / 2
        y = A4[1] - 10*mm - logo_h   # 10mm from top edge
        canvas.drawImage(LOGO_PATH, x, y, width=logo_w, height=logo_h, mask='auto')
    canvas.setFont('Helvetica', 10)
    canvas.setFillColor('#808080')
    canvas.drawCentredString(A4[0]/2, 12*mm, 'Rua Pará, 945   |   Centro   |   Londrina/PR   |   CEP 86010-450   |   @genesisregistro')
    canvas.restoreState()

# ── Markdown-to-RML helpers ───────────────────────────────────────────
def md_bold(text):
    """Convert **bold** markdown to <b> tags."""
    parts = text.split('**')
    out = ""
    for i, p in enumerate(parts):
        out += f"<b>{p}</b>" if i % 2 else p
    return out

def _is_no_indent(line):
    """Lines that should NOT have first-line indent."""
    up = line.upper()
    triggers = ['CONTRATANTE:', 'CONTRATADO:', 'OUTORGANTE', 'OUTORGADO',
                '[NO_INDENT]', 'LONDRINA,']
    return any(t in up for t in triggers)

# ── PDF Builder ────────────────────────────────────────────────────────
def build_pdf(content, filename, styles, is_proc=False):
    top_m = 3.5*cm if is_proc else 4.5*cm
    doc = SimpleDocTemplate(filename, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=top_m, bottomMargin=2.5*cm)
    els = []
    norm  = styles['PN']  if is_proc else styles['GN']
    noi   = styles['PNI'] if is_proc else styles['GNI']
    in_sig = False
    sig_elements = []  # collect signature block for KeepTogether

    for line in content.split('\n'):
        s = line.strip()
        if not s:
            if in_sig:
                sig_elements.append(Spacer(1, 2*mm))
            else:
                els.append(Spacer(1, 2*mm))
            continue

        # Detect start of signature zone: line that starts with "Londrina," (the closing line)
        # Pull the last paragraph (+ preceding spacers) so signatures never appear alone
        if not in_sig and re.match(r'^Londrina,\s', s, re.IGNORECASE):
            in_sig = True
            # Pull back last paragraph and its spacers from els into sig_elements
            pulled = []
            while els and isinstance(els[-1], Spacer):
                pulled.insert(0, els.pop())
            if els and isinstance(els[-1], Paragraph):
                pulled.insert(0, els.pop())
                # Also grab spacers before that paragraph
                while els and isinstance(els[-1], Spacer):
                    pulled.insert(0, els.pop())
            sig_elements.extend(pulled)
            clean = s.replace('[NO_INDENT]', '')
            style = noi if _is_no_indent(s) else norm
            sig_elements.append(Paragraph(clean, style))
            continue

        # Signatures
        if s.startswith('__'):
            if not in_sig:
                in_sig = True
            sig_elements.append(Spacer(1, 8*mm))
            sig_elements.append(Paragraph(s, styles['GSIG']))
            continue
        if in_sig:
            sig_elements.append(Paragraph(s, styles['GSIG']))
            continue

        # Headers
        if s.startswith('# '):
            if not is_proc:
                els.append(Paragraph(s[2:], styles['GH1']))
            continue
        if s.startswith('## '):
            els.append(Paragraph(s[3:], styles['GH2']))
            continue

        # Regular paragraph
        clean = s.replace('[NO_INDENT]', '')
        style = noi if _is_no_indent(s) else norm
        els.append(Paragraph(clean, style))

    # Wrap entire signature block in KeepTogether so it never splits across pages
    if sig_elements:
        els.append(KeepTogether(sig_elements))

    doc.build(els, onFirstPage=_header_footer, onLaterPages=_header_footer)
    print(f"Generated PDF: {filename}")

# ── Qualification Generators ──────────────────────────────────────────
GENESIS = {
    "razao_social": "LUAN VIEIRA MENDES LTDA",
    "nome_rep": "Luan Vieira Mendes",
    "cnpj": "51.829.412/0001-70",
    "endereco_sede": "Rua Joao Wyclif, 405 - Gleba Fazenda Palhano, Londrina",
    "cargo_rep": "sócio-administrador",
    "nacionalidade_rep": "brasileiro",
    "estado_civil_rep": "empresário",
    "profissao_rep": "empresário",
    "cpf_rep": "085.529.039-03",
    "rg_rep": "N/A",
    "orgao_emissor_rep": "",
    "endereco_rep": "Rua Joao Wyclif, 405, Londrina",
}

def qual_pj(d):
    r = f"<b><u>{d['razao_social']}</u></b>"
    n = f"<b><u>{d['nome_rep']}</u></b>"
    return (f"[NO_INDENT]{r}, pessoa jurídica de direito privado, inscrita no CNPJ "
            f"sob o nº {d['cnpj']}, com sede na {d['endereco_sede']}, "
            f"CEP {d.get('cep_sede','N/A')}, neste ato representada por seu "
            f"{d['cargo_rep']}, {n}, {d['nacionalidade_rep']}, "
            f"{d['estado_civil_rep']}, {d['profissao_rep']}, portador do RG nº "
            f"{d['rg_rep']} {d.get('orgao_emissor_rep','')}, inscrito no CPF sob "
            f"o nº {d['cpf_rep']}, residente e domiciliado na {d['endereco_rep']}.")

def qual_pf(d):
    n = f"<b><u>{d['nome']}</u></b>"
    return (f"[NO_INDENT]{n}, {d['nacionalidade']}, {d['estado_civil']}, "
            f"{d['profissao']}, portador do RG nº {d['rg']} "
            f"{d.get('orgao_emissor','')}, inscrito no CPF sob o nº {d['cpf']}, "
            f"residente e domiciliado na {d['endereco']}, CEP {d.get('cep','N/A')}.")

# ── Placeholder Engine ────────────────────────────────────────────────
def replace_all(text, data):
    """Replace {{KEY}} placeholders, then convert ** to <b>, then strip leftovers."""
    result = text
    # 1. Replace all data keys (case-insensitive)
    for k, v in data.items():
        if isinstance(v, str):
            result = re.sub(re.escape('{{'+k+'}}'), v, result, flags=re.IGNORECASE)
            result = re.sub(re.escape('{{'+k.upper()+'}}'), v, result, flags=re.IGNORECASE)
            result = re.sub(re.escape('{{'+k.lower()+'}}'), v, result, flags=re.IGNORECASE)
    # 2. Cleanup any remaining placeholders → empty string
    result = re.sub(r'\{\{[A-Za-z0-9_]+\}\}', '', result)
    # 3. Convert ** to <b>
    result = md_bold(result)
    return result

# ── DOCX Builder ──────────────────────────────────────────────────────
def _docx_add_logo_header(doc):
    """Add Genesis logo to header (all sections)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(LOGO_PATH):
            run = p.add_run()
            run.add_picture(LOGO_PATH, width=Cm(4))

def _docx_add_footer(doc):
    """Add Genesis footer to all sections."""
    footer_text = 'Rua Pará, 945   |   Centro   |   Londrina/PR   |   CEP 86010-450   |   @genesisregistro'
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(footer_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        _docx_set_font(run, 'Sora')

def _docx_set_font(run, font_name='Sora'):
    """Set font on a run, ensuring Word picks it up correctly."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def _docx_add_bold_text(paragraph, text, font_name='Sora', font_size=12, color=None, underline=False):
    """Parse **bold** markdown and add runs to paragraph."""
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(font_size)
        run.bold = (i % 2 == 1)
        if underline and (i % 2 == 1):
            run.underline = True
        if color:
            run.font.color.rgb = color
        _docx_set_font(run, font_name)

def _docx_strip_html(text):
    """Strip HTML tags (<b>, <u>, etc.) and convert to **bold** markdown for DOCX."""
    text = re.sub(r'<b><u>(.*?)</u></b>', r'**\1**', text)
    text = re.sub(r'<b>(.*?)</b>', r'**\1**', text)
    text = re.sub(r'<u>(.*?)</u>', r'**\1**', text)
    text = text.replace('[NO_INDENT]', '')
    return text

def build_docx(content, filename, is_proc=False):
    """Build a DOCX document from the same markdown content used for PDF."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(4.5) if not is_proc else Cm(3.5)
    section.bottom_margin = Cm(2.5)

    # Header & footer
    _docx_add_logo_header(doc)
    _docx_add_footer(doc)

    font_size = 11 if is_proc else 12
    line_spacing = 1.5

    # Check if Sora font files exist (DOCX embeds by name, user needs font installed)
    sora_r = os.path.join(ASSETS_DIR, 'Sora-Regular.ttf')
    font_name = 'Sora' if os.path.exists(sora_r) else 'Calibri'

    for line in content.split('\n'):
        s = line.strip()

        if not s:
            doc.add_paragraph()
            continue

        # Strip HTML tags from content (qualifications use <b><u>)
        s = _docx_strip_html(s)

        # Headers
        if s.startswith('# '):
            if not is_proc:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(s[2:])
                run.bold = True
                run.font.size = Pt(16)
                _docx_set_font(run, font_name)
            continue

        if s.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(s[3:])
            run.bold = True
            run.font.size = Pt(12)
            _docx_set_font(run, font_name)
            continue

        # Signature lines
        if s.startswith('__'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            _docx_add_bold_text(p, s, font_name, 11)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.space_after = Pt(6 if not is_proc else 4)

        # First-line indent (like PDF)
        if not _is_no_indent(s) and not s.startswith('- '):
            p.paragraph_format.first_line_indent = Cm(2)

        _docx_add_bold_text(p, s, font_name, font_size)

    doc.save(filename)
    print(f"Generated DOCX: {filename}")

# ── Main ───────────────────────────────────────────────────────────────
def main():
    if len(sys.argv) < 2:
        print("Usage: python3 assemble_contract.py '<json>'")
        sys.exit(1)
    data = json.loads(sys.argv[1])

    # Date
    M = {1:'janeiro',2:'fevereiro',3:'março',4:'abril',5:'maio',6:'junho',
         7:'julho',8:'agosto',9:'setembro',10:'outubro',11:'novembro',12:'dezembro'}
    now = datetime.datetime.now()
    data['DATA_ATUAL'] = f"{now.day} de {M[now.month]} de {now.year}"

    # Qualifications
    q_contratado = qual_pj(GENESIS)
    q_contratante = qual_pj(data) if data.get('tipo_pessoa')=='PJ' else qual_pf(data)

    # Clause selection
    cl_classes  = CLAUSES['classes']['especifica' if data.get('lista_classes') else 'quantidade']
    cl_pagto    = CLAUSES['pagamento'].get(data.get('forma_pagamento','pix_manual'), CLAUSES['pagamento']['pix_manual'])
    cl_cond     = CLAUSES['condicoes'].get(data.get('condicao_pagamento','vista'), CLAUSES['condicoes']['vista'])

    # Signature block
    if data.get('tipo_pessoa') == 'PJ':
        sig = f"_____________________________________\n{data.get('nome_rep','')}\nRepresentante Legal de {data.get('razao_social','')}\nContratante"
    else:
        sig = f"_____________________________________\n{data.get('nome','')}\nContratante"

    # Assembly helper
    def assemble(tpl_file):
        path = os.path.join(TERMOS_DIR, tpl_file)
        with open(path, 'r', encoding='utf-8') as f:
            t = f.read()
        t = t.replace('{{QUALIFICACAO_CONTRATANTE}}', q_contratante)
        t = t.replace('{{QUALIFICACAO_CONTRATADO}}', q_contratado)
        t = t.replace('{{CLAUSULA_CLASSES}}', cl_classes)
        t = t.replace('{{CLAUSULA_PAGAMENTO}}', cl_pagto)
        t = t.replace('{{CLAUSULA_CONDICOES}}', cl_cond)
        t = t.replace('{{ASSINATURA_CONTRATANTE}}', sig)
        return replace_all(t, data)

    fn, fb = setup_fonts()
    styles = build_styles(fn, fb)
    sfx = data.get('nome_arquivo_suffix', 'Cliente')
    generated_files = []

    if data.get('gerar_contrato', True):
        content = assemble('contrato.md')
        fname_pdf = f"Contrato_{sfx}.pdf"
        fname_docx = f"Contrato_{sfx}.docx"
        build_pdf(content, fname_pdf, styles, is_proc=False)
        build_docx(content, fname_docx, is_proc=False)
        generated_files.append(fname_pdf)
        generated_files.append(fname_docx)

    if data.get('gerar_procuracao', True):
        content = assemble('procuracao.md')
        fname_pdf = f"Procuracao_{sfx}.pdf"
        fname_docx = f"Procuracao_{sfx}.docx"
        build_pdf(content, fname_pdf, styles, is_proc=True)
        build_docx(content, fname_docx, is_proc=True)
        generated_files.append(fname_pdf)
        generated_files.append(fname_docx)

    # ── Google Drive Upload ──────────────────────────────────────────
    if data.get('upload_drive', True):
        try:
            from google_drive_service import upload_file
            marca = data.get('nome_marca', 'Sem Marca')
            cliente = data.get('razao_social') if data.get('tipo_pessoa') == 'PJ' else data.get('nome')
            
            for fpath in generated_files:
                upload_file(os.path.abspath(fpath), marca, cliente)
        except Exception as e:
            print(f"[ERROR] Google Drive upload failed: {e}")

if __name__ == "__main__":
    main()
